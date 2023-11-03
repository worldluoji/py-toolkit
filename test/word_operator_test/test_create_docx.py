import os
from docx import Document
from docx.shared import Inches

TEST_FILE_NAME = 'test.docx'
def test_create_docx():
    document = Document()

    document.add_heading('Document Title', 0)

    # create and add a new paragraph
    p = document.add_paragraph('A plain paragraph having some ')
    # add style to this paragraph
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('/Users/honorluo/Downloads/test/WechatIMG81.jpeg', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save(TEST_FILE_NAME)
    assert os.path.exists(os.path.join(os.getcwd(), TEST_FILE_NAME)) == True

    os.remove(TEST_FILE_NAME)
     


    