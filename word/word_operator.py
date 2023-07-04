from docx import Document
from docx.shared import RGBColor


def add_to_end(doc, content, font='宋体', underline = False, color='', save_to= ''):
    para = doc.add_paragraph().add_run(content)
    # set font style
    para.font.name = font
    # set underline
    para.font.underline = underline
    # set color
    if color != '':
        para.font.color.rgb = color # RGBColor(255,128,128)
    
    if (save_to is not None) and len(save_to) > 0:
        try:
            doc.save(save_to)
        except Exception as e:
            print(e)


def add_before_text(doc, text, content, stop = True, save_to= ''):
    for para in doc.paragraphs:
        if para.text == text:
            para.insert_paragraph_before(content)
            if stop:
                break
    
    if (save_to is not None) and len(save_to) > 0:
        try:
            doc.save(save_to)
        except Exception as e:
            print(e)