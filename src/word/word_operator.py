from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from enum import Enum  

# https://python-docx.readthedocs.io/en/latest/#

class RETUENED_STATUS(Enum):  
    SUCCESS = 0
    NOT_CHANGED = 1
    FAIL_TO_SAVE = 2


'''
替换段落中的字符串,做法整体是两种:
1. 对paragraph中的文字进行替换(replace_para)。但是这有一个问题,原来整段的文字格式都会丢失。
2. 遍历paragraph.runs中,对其中的每一段文字进行判断和替换。有个漏洞,就是需要被替换的字符串可能会被拆分到多个run中,导致匹配不到

相对比较好的解决办法：
对runs中的内容进行一定程度的拼接,但是有缺点,部分文字的样式可能会消失,可以尽量让每一段文字的样式保持一致来避免这种情况。

:param doc: the word document to open
:param keyword: the keyword to be replaced
:param content: the content to replace the keyword
:param save_to: indicates where to save the modified doc
:return RETUENED_STATUS
'''
def replace_paragraph_text(doc, keyword, content, save_to= ''):
    for paragraph in doc.paragraphs:
        if keyword not in paragraph.text:
            continue
        tmp = ''
        runs = paragraph.runs
        for i, run in enumerate(runs):
            tmp += run.text  # 合并run字符串
            if keyword in tmp:
                # 如果存在匹配得字符串,那么将当前得run替换成合并后得字符串
                run.text = run.text.replace(run.text, tmp)
                run.text = run.text.replace(keyword, content)
                tmp = ''
            else:
                # 如果没匹配到目标字符串则把当前run置空
                run.text = run.text.replace(run.text, '')
            if i == len(runs) - 1:
                # 如果是当前段落一直没有符合规则得字符串直接将当前run替换为tmp
                run.add_text(tmp)

    return save(doc, save_to)


'''替换table某个cell中的text
:param doc: the word document to open
:param keyword: the keyword to be replaced
:param content: the content to replace the keyword
:param save_to: indicates where to save the modified doc
:return RETUENED_STATUS
'''
def replace_table_cell_text(doc, keyword, content, save_to= ''):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 如果只是为了内容,直接替换cell.text,但是为了保存原有格式,需要将每个单元格的文本当作一段看待,以此提取出run来不修改原格式
                for paragraph in cell.paragraphs:
                    if keyword in paragraph.text:
                        has_replaced = False
                        for run in paragraph.runs:
                            run.clear()
                            if not has_replaced:
                                run.add_text(content)
                                has_replaced = True

    return save(doc, save_to)



default_styles = {'font': {'name': '宋体', 'color': RGBColor(0,0,0), 'bold': False, 'underline': False, 'size': Pt(12)},
'paragraph_format': {'alignment': WD_ALIGN_PARAGRAPH.LEFT, 'left_indent': Pt(2)}}

'''set the paragraph font styles
    :param para: the paragraph style object to be set
    :param styles: a dict accroding to the reference: https://python-docx.readthedocs.io/en/latest/api/style.html#docx.styles.style.ParagraphStyle
'''
def set_font_styles(para, styles):
    if styles['font'] is None:
        return

    para.font.name = styles['font']['name'] if styles['font']['name'] is not None else default_styles['font']['name']
    para.font.color.rgb = styles['font']['color'] if styles['font']['color'] is not None else default_styles['font']['color']
    para.font.bold = styles['font']['bold'] if styles['font']['bold'] is not None else default_styles['font']['bold']
    para.font.underline = styles['font']['underline'] if styles['font']['underline'] is not None else default_styles['font']['underline']
    para.font.size = styles['font']['size'] if styles['font']['size'] is not None else default_styles['font']['size']

'''add content in end of the doc
   :param doc: the word document to open
   :param content: the content to be insert
   :param font: the content's font style
   :param underline: if content with underline
   :param save_to: indicates where to save the modified doc
'''
def add_to_end(doc, content, save_to= '', styles=default_styles):
    para = doc.add_paragraph().add_run(content)
    set_font_styles(para, styles)
    if (save_to is not None) and len(save_to) > 0:
        try:
            doc.save(save_to)
        except Exception as e:
            print(e)
            return RETUENED_STATUS.FAIL_TO_SAVE.value

    return RETUENED_STATUS.SUCCESS.value

'''add content before text in doc
   :param doc: the word document to open
   :param keyword: the content to be insert before keyword
   :param content: the content to be insert
   :param stop: when found the first keyword then break if stop = True
   :param save_to: indicates where to save the modified doc
   :param styles: a dict accroding to the reference: https://python-docx.readthedocs.io/en/latest/api/style.html#docx.styles.style.ParagraphStyle
   :return RETUENED_STATUS
'''
def add_before_text(doc, keyword, content, stop = True, save_to= '', styles = default_styles):
    ret = RETUENED_STATUS.NOT_CHANGED.value
    for para in doc.paragraphs:
        if para.text == keyword:
            newPara = para.insert_paragraph_before().add_run(content)
            ret = RETUENED_STATUS.SUCCESS.value
            set_font_styles(newPara, styles) 
            if stop:
                break
    
    if (save_to is not None) and len(save_to) > 0:
        try:
            doc.save(save_to)
        except Exception as e:
            print(e)
            ret = RETUENED_STATUS.FAIL_TO_SAVE.value
    
    return ret


'''replace paragraph's content which content is equal to text
   :param doc: the word document to open
   :param keyword: the content to be insert before keyword
   :param content: the content to be insert
   :param stop: when found the first keyword then break if stop = True
   :param save_to: indicates where to save the modified doc
   :param styles: a dict accroding to the reference: https://python-docx.readthedocs.io/en/latest/api/style.html#docx.styles.style.ParagraphStyle
   :return RETUENED_STATUS
'''
def replace_para(doc, keyword, content, stop = True, save_to= '', styles = default_styles):
    ret = RETUENED_STATUS.NOT_CHANGED.value
    for para in doc.paragraphs:
        if para.text == keyword:
            para.text = content
            ret = RETUENED_STATUS.SUCCESS.value
            set_font_styles(para.style, styles)
            if stop:
                break
        
    if (save_to is not None) and len(save_to) > 0:
        try:
            doc.save(save_to)
        except Exception as e:
            print(e)
            ret = RETUENED_STATUS.FAIL_TO_SAVE.value

    return ret


def save(doc, save_to= ''):
    if (save_to is not None) and len(save_to) > 0:
        try:
            doc.save(save_to)
        except Exception as e:
            print(e)
            return RETUENED_STATUS.FAIL_TO_SAVE.value

    return RETUENED_STATUS.SUCCESS.value